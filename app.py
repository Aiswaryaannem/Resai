import streamlit as st
import requests
from docx import Document

# Configure DeepSeek API Key
DEEPSEEK_API_KEY = 'your key'  # Replace with your actual API key
DEEPSEEK_API_URL = 'https://api.deepseek.com/v1/chat/completions'

if not DEEPSEEK_API_KEY:
    st.error("API key is missing. Please check your configuration.")

# Function to extract text from DOCX
def extract_docx_content(docx_file):
    """
    Extracts text from a DOCX file while preserving the original formatting.
    """
    try:
        doc = Document(docx_file)
        sections = [{"text": para.text, "style": para.style} for para in doc.paragraphs]
        return sections
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None

# Function to enhance resume content
def enhance_resume_content(current_sections, job_description):
    """
    Enhances resume content section by section based on the job description.
    """
    try:
        current_text = "\n".join([section["text"] for section in current_sections if section["text"].strip()])
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        data = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "system",
                    "content": (
                        "You are a professional resume writer and ATS optimization expert. "
                        "Rewrite the provided resume to align perfectly with the job description. "
                        "Ensure it includes relevant skills, keywords, and experiences while retaining its original format. "
                        "The resume must be ATS-friendly and free of gaps."
                    ),
                },
                {
                    "role": "user",
                    "content": f"Resume:\n{current_text}\n\nJob Description:\n{job_description}",
                },
            ],
            "max_tokens": 3000,
        }
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        response.raise_for_status()  # Raise an error for bad status codes
        return response.json()["choices"][0]["message"]["content"]
    except requests.exceptions.HTTPError as err:
        if response.status_code == 402:
            st.error("Payment Required: Please subscribe to a plan or check your payment status.")
        elif response.status_code == 401:
            st.error("Unauthorized: Invalid API key. Please check your API key.")
        else:
            st.error(f"HTTP Error: {err}")
    except Exception as e:
        st.error(f"Error enhancing resume content: {e}")
    return None

# Function to save enhanced content as Word Document
def save_word_content(updated_content):
    """
    Saves the enhanced resume content into a Word document.
    """
    try:
        doc = Document()
        # Split the content into lines and add each line as a paragraph
        lines = updated_content.split("\n")
        for line in lines:
            if line.strip():  # Only add non-empty lines
                doc.add_paragraph(line)
        
        filename = "Enhanced_Resume.docx"
        doc.save(filename)
        return filename
    except Exception as e:
        st.error(f"Error saving enhanced resume as Word document: {e}")
        return None

# Streamlit app interface
st.title("RESAI")
st.text("Upload your DOCX resume and paste the job description to generate a clean, ATS-optimized Word resume.")

# Upload DOCX resume file
uploaded_file = st.file_uploader("Upload Your Resume (DOCX)", type=["docx"], help="Upload your Word resume in DOCX format.")

# Input job description
job_description = st.text_area("Paste the Job Description", help="Enter the job description for the role you are applying for.")

# Submit button
if st.button("Enhance Resume"):
    if not uploaded_file:
        st.warning("Please upload your current resume.")
    elif not job_description:
        st.warning("Please paste the job description.")
    else:
        st.info("Extracting content from the uploaded resume...")
        current_sections = extract_docx_content(uploaded_file)

        if current_sections:
            st.success("Resume content extracted successfully.")
            st.info("Enhancing resume content based on the job description...")
            enhanced_content = enhance_resume_content(current_sections, job_description)

            if enhanced_content:
                st.success("Resume enhanced successfully!")
                st.info("Saving enhanced resume as Word document...")
                enhanced_word_file = save_word_content(enhanced_content)

                if enhanced_word_file:
                    st.success("Enhanced resume saved as Word document.")
                    with open(enhanced_word_file, "rb") as file:
                        st.download_button(
                            "Download Enhanced Resume (Word)",
                            file,
                            file_name="Enhanced_Resume.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                else:
                    st.error("Failed to save enhanced resume as Word document. Please try again.")
            else:
                st.error("Failed to enhance resume content. Please try again.")
        else:
            st.error("Failed to extract content from the resume. Please ensure the file is a valid DOCX format.")
